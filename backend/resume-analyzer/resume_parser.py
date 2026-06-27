import os
import zipfile
import xml.etree.ElementTree as ET
import pypdf

try:
    import docx
except ImportError:
    docx = None

class ResumeParser:
    """
    Parses resume files (PDF, TXT, DOCX) and extracts raw text.
    """
    @staticmethod
    def parse_pdf(file_path: str) -> str:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        text = []
        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        return "\n".join(text).strip()

    @staticmethod
    def parse_txt(file_path: str) -> str:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()

    @staticmethod
    def parse_docx(file_path: str) -> str:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        # Prefer python-docx when installed
        if docx is not None:
            document = docx.Document(file_path)
            parts = [para.text for para in document.paragraphs if para.text]
            for table in document.tables:
                for row in table.rows:
                    row_text = " ".join(cell.text for cell in row.cells if cell.text)
                    if row_text:
                        parts.append(row_text)
            return "\n".join(parts).strip()

        # Fallback for DOCX parsing without python-docx
        if not zipfile.is_zipfile(file_path):
            raise ValueError("DOCX file is not a valid ZIP package.")

        with zipfile.ZipFile(file_path) as archive:
            document_path = None
            for name in archive.namelist():
                if name.lower().endswith("word/document.xml"):
                    document_path = name
                    break
            if document_path is None:
                raise ValueError("Invalid DOCX file: missing word/document.xml entry.")
            xml_content = archive.read(document_path)

        root = ET.fromstring(xml_content)
        paragraphs = []
        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        for paragraph in root.findall('.//w:p', namespace):
            texts = [node.text for node in paragraph.findall('.//w:t', namespace) if node.text]
            if texts:
                paragraphs.append("".join(texts))
        return "\n".join(paragraphs).strip()

    @classmethod
    def parse(cls, file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return cls.parse_pdf(file_path)
        elif ext in (".txt", ".md"):
            return cls.parse_txt(file_path)
        elif ext == ".docx":
            return cls.parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format '{ext}'. Only PDF, TXT, and DOCX are supported.")
